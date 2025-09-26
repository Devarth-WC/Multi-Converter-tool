from gtts import gTTS
import pyttsx3
from docx import Document
import os
import tempfile

class TextConverter:
    def __init__(self):
        pass
    
    def text_to_audio(self, text_file_path, output_path, engine='gtts'):
        """Convert text file to audio"""
        try:
            # Read text content
            with open(text_file_path, 'r', encoding='utf-8') as file:
                text_content = file.read()
            
            if not text_content.strip():
                raise Exception("No text content found in file")
            
            if engine == 'gtts':
                # Use Google Text-to-Speech (requires internet)
                self._text_to_audio_gtts(text_content, output_path)
            elif engine == 'pyttsx3':
                # Use offline text-to-speech
                self._text_to_audio_offline(text_content, output_path)
            else:
                raise Exception(f"Unsupported TTS engine: {engine}")
            
            return output_path
            
        except Exception as e:
            raise Exception(f"Text to audio conversion failed: {str(e)}")
    
    def _text_to_audio_gtts(self, text, output_path):
        """Convert text to audio using Google TTS"""
        # Limit text length for gTTS
        if len(text) > 4500:
            text = text[:4500] + "..."
        
        tts = gTTS(text=text, lang='en', slow=False)
        tts.save(output_path)
    
    def _text_to_audio_offline(self, text, output_path):
        """Convert text to audio using offline TTS"""
        engine = pyttsx3.init()
        
        # Configure voice properties
        voices = engine.getProperty('voices')
        if voices:
            engine.setProperty('voice', voices[0].id)  # Use first available voice
        
        engine.setProperty('rate', 150)    # Speed of speech
        engine.setProperty('volume', 0.9)  # Volume level (0.0 to 1.0)
        
        # Save to file
        engine.save_to_file(text, output_path)
        engine.runAndWait()
    
    def docx_to_txt(self, docx_path, output_path):
        """Convert DOCX to plain text"""
        try:
            doc = Document(docx_path)
            text_content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + "\t"
                    text_content += "\n"
            
            # Save as text file
            with open(output_path, 'w', encoding='utf-8') as txt_file:
                txt_file.write(text_content)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"DOCX to TXT conversion failed: {str(e)}")
    
    def txt_to_docx(self, txt_path, output_path):
        """Convert plain text to DOCX"""
        try:
            # Read text content
            with open(txt_path, 'r', encoding='utf-8') as txt_file:
                text_content = txt_file.read()
            
            # Create new document
            doc = Document()
            
            # Add paragraphs (split by double newlines)
            paragraphs = text_content.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            
            # Save document
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"TXT to DOCX conversion failed: {str(e)}")
    
    def format_text(self, text_path, output_path, formatting_options=None):
        """Apply formatting to text file"""
        try:
            if formatting_options is None:
                formatting_options = {}
            
            with open(text_path, 'r', encoding='utf-8') as file:
                text_content = file.read()
            
            # Apply formatting options
            if formatting_options.get('uppercase'):
                text_content = text_content.upper()
            elif formatting_options.get('lowercase'):
                text_content = text_content.lower()
            elif formatting_options.get('title_case'):
                text_content = text_content.title()
            
            if formatting_options.get('remove_extra_spaces'):
                lines = text_content.split('\n')
                text_content = '\n'.join(line.strip() for line in lines)
            
            if formatting_options.get('add_line_numbers'):
                lines = text_content.split('\n')
                numbered_lines = []
                for i, line in enumerate(lines, 1):
                    numbered_lines.append(f"{i:3d}: {line}")
                text_content = '\n'.join(numbered_lines)
            
            # Save formatted text
            with open(output_path, 'w', encoding='utf-8') as file:
                file.write(text_content)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"Text formatting failed: {str(e)}")
    
    def merge_text_files(self, text_paths, output_path, separator="\n\n---\n\n"):
        """Merge multiple text files into one"""
        try:
            merged_content = ""
            
            for i, text_path in enumerate(text_paths):
                with open(text_path, 'r', encoding='utf-8') as file:
                    content = file.read()
                    merged_content += content
                    
                    # Add separator between files (except for the last file)
                    if i < len(text_paths) - 1:
                        merged_content += separator
            
            # Save merged content
            with open(output_path, 'w', encoding='utf-8') as file:
                file.write(merged_content)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"Text file merging failed: {str(e)}")
    
    def extract_keywords(self, text_path, output_path, min_length=4):
        """Extract keywords from text file"""
        try:
            with open(text_path, 'r', encoding='utf-8') as file:
                text_content = file.read()
            
            # Simple keyword extraction (can be enhanced with NLP libraries)
            words = text_content.lower().split()
            
            # Filter words by length and remove common stop words
            stop_words = {'the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by'}
            keywords = []
            
            for word in words:
                # Remove punctuation
                clean_word = ''.join(char for char in word if char.isalnum())
                if len(clean_word) >= min_length and clean_word not in stop_words:
                    keywords.append(clean_word)
            
            # Count frequency and get unique keywords
            from collections import Counter
            word_counts = Counter(keywords)
            
            # Save keywords with frequency
            with open(output_path, 'w', encoding='utf-8') as file:
                file.write("Keywords (with frequency):\n\n")
                for word, count in word_counts.most_common(50):  # Top 50 keywords
                    file.write(f"{word}: {count}\n")
            
            return output_path
            
        except Exception as e:
            raise Exception(f"Keyword extraction failed: {str(e)}")